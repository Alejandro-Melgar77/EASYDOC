from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

REFERENCE = Path(r"C:\Users\Usuario\Documents\SW1 1-2026\modelos\para el segundo parcial\documento 1--2026\Sistema de gestion de tramites y politicas de negocio (1).docx")
OUT = Path(r"C:\Users\Usuario\Documents\SW1 1-2026\segundo parcial\puds easydoc\qa\modelo_inventario.txt")

def p_style(paragraph):
    ppr = paragraph._p.pPr
    if ppr is None:
        return ""
    pstyle = ppr.find(qn("w:pStyle"))
    return pstyle.get(qn("w:val")) if pstyle is not None else ""

def walk_table(table, depth=0):
    rows = []
    indent = "  " * depth
    for r_idx, row in enumerate(table.rows, 1):
        cells = []
        for c_idx, cell in enumerate(row.cells, 1):
            text = " | ".join(p.text.strip().replace("\n", " ") for p in cell.paragraphs if p.text.strip())
            cells.append(text)
        rows.append(f"{indent}  fila {r_idx}: " + " || ".join(cells))
    return rows

doc = Document(str(REFERENCE))
lines = []
lines.append(f"PARAGRAPH_COUNT: {len(doc.paragraphs)}")
lines.append(f"TABLE_COUNT: {len(doc.tables)}")
lines.append("\n=== PARRAFOS ===")
for i, p in enumerate(doc.paragraphs, 1):
    text = p.text.strip().replace("\n", " ")
    if text:
        lines.append(f"{i:04d} [style={p.style.name!r}; style_id={p_style(p)!r}; align={p.alignment}] {text}")
lines.append("\n=== TABLAS ===")
for i, table in enumerate(doc.tables, 1):
    lines.append(f"\nTABLA {i}: filas={len(table.rows)}, columnas={len(table.columns)}, style={table.style.name if table.style else ''}")
    lines.extend(walk_table(table))

OUT.write_text("\n".join(lines), encoding="utf-8")
print(OUT)
