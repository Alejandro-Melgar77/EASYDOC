from __future__ import annotations

import copy
import html
import shutil
import textwrap
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Usuario\Documents\SW1 1-2026\segundo parcial")
OUT = ROOT / "puds easydoc"
ASSETS = OUT / "activos"
QA = OUT / "qa"
REFERENCE = Path(r"C:\Users\Usuario\Documents\SW1 1-2026\modelos\para el segundo parcial\documento 1--2026\Sistema de gestion de tramites y politicas de negocio (1).docx")
FINAL = OUT / "documento easydoc.docx"
DRAWIO = ASSETS / "diagramas_easydoc.drawio"

NAVY = "#163A5F"
TEAL = "#0F766E"
GOLD = "#B88924"
INK = "#1F2937"
PALE = "#EAF2F8"
GREEN = "#1F8A70"
RED = "#B42318"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


FONT_20 = font(20, True)
FONT_18 = font(18, True)
FONT_16 = font(16, True)
FONT_14 = font(14)
FONT_13 = font(13)
FONT_12 = font(12)


def text_center(draw, box, text, fnt=FONT_14, fill=INK, max_chars=25):
    x1, y1, x2, y2 = box
    lines = []
    for piece in str(text).split("\n"):
        lines.extend(textwrap.wrap(piece, width=max_chars) or [""])
    bboxes = [draw.textbbox((0, 0), line, font=fnt) for line in lines]
    line_h = max((b[3] - b[1] for b in bboxes), default=14) + 4
    total = line_h * len(lines)
    y = y1 + max(5, ((y2 - y1) - total) / 2)
    for line, bb in zip(lines, bboxes):
        w = bb[2] - bb[0]
        draw.text(((x1 + x2 - w) / 2, y), line, font=fnt, fill=fill)
        y += line_h


def rounded(draw, box, fill, outline=NAVY, radius=18, width=3, text=None, fnt=FONT_14, text_fill=INK):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    if text:
        text_center(draw, box, text, fnt=fnt, fill=text_fill)


def ellipse(draw, box, fill, outline=NAVY, width=3, text=None, fnt=FONT_14):
    draw.ellipse(box, fill=fill, outline=outline, width=width)
    if text:
        text_center(draw, box, text, fnt=fnt)


def arrow(draw, start, end, fill=NAVY, width=4, label=None):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        points = [(x2, y2), (x2 - 15 if x2 > x1 else x2 + 15, y2 - 8), (x2 - 15 if x2 > x1 else x2 + 15, y2 + 8)]
    else:
        points = [(x2, y2), (x2 - 8, y2 - 15 if y2 > y1 else y2 + 15), (x2 + 8, y2 - 15 if y2 > y1 else y2 + 15)]
    draw.polygon(points, fill=fill)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.text((mx + 6, my - 20), label, font=FONT_12, fill=INK)


def title(draw, title_text, subtitle):
    draw.rectangle((0, 0, 1600, 78), fill=NAVY)
    draw.text((45, 18), title_text, font=FONT_20, fill=WHITE)
    draw.text((45, 92), subtitle, font=FONT_13, fill=INK)


def save_img(img, name):
    path = ASSETS / name
    img.save(path, quality=95)
    return path


def diag_use_cases():
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title(d, "EASYDOC - Diagrama general de casos de uso", "Actores institucionales y módulos de la plataforma académica")
    d.rectangle((275, 155, 1325, 830), outline=NAVY, width=4, fill="#F8FBFE")
    d.text((305, 170), "EASYDOC", font=FONT_18, fill=NAVY)
    actors = [(80, 245, "Director /\nAdministrador"), (80, 515, "Funcionario /\nSecretaría"), (1400, 370, "Estudiante /\nUsuario invitado"), (1400, 610, "Servicios IA\nlocales")]
    for x, y, name in actors:
        d.ellipse((x + 35, y, x + 75, y + 40), outline=INK, width=3)
        d.line((x + 55, y + 40, x + 55, y + 105), fill=INK, width=3)
        d.line((x + 20, y + 65, x + 90, y + 65), fill=INK, width=3)
        d.line((x + 55, y + 105, x + 25, y + 150), fill=INK, width=3)
        d.line((x + 55, y + 105, x + 85, y + 150), fill=INK, width=3)
        text_center(d, (x - 10, y + 160, x + 120, y + 205), name, FONT_13)
    cases = [
        (360, 250, "Gestionar usuarios, roles\ny gobierno"),
        (690, 250, "Diseñar políticas UML\ny formularios"),
        (1015, 250, "Publicar y versionar\npolíticas"),
        (360, 450, "Registrar y avanzar\ntrámites"),
        (690, 450, "Gestionar expedientes\ny colaboración"),
        (1015, 450, "Consultar mesa de control\ny trazabilidad"),
        (520, 650, "Asistente EASY\ny reportes"),
        (865, 650, "Predecir riesgos\noperativos"),
    ]
    for x, y, label in cases:
        ellipse(d, (x, y, x + 250, y + 90), "#DFF3F0", TEAL, text=label, fnt=FONT_14)
    for end in [(360, 295), (690, 295), (1015, 295), (1015, 495), (1015, 695)]:
        arrow(d, (190, 330), end, fill="#5B7083", width=2)
    for end in [(360, 495), (690, 495), (520, 695)]:
        arrow(d, (190, 600), end, fill="#5B7083", width=2)
    for end in [(690, 495), (520, 695)]:
        arrow(d, (1400, 475), end, fill="#5B7083", width=2)
    for end in [(520, 695), (865, 695)]:
        arrow(d, (1400, 700), end, fill="#5B7083", width=2)
    return save_img(img, "01_diagrama_casos_uso_easydoc.png")


def diag_activity():
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title(d, "EASYDOC - Flujo de actividad: casos especiales académicos", "Política publicada y registrada en la biblioteca de políticas")
    d.rectangle((70, 155, 520, 820), outline="#9FB3C8", width=2, fill="#F8FBFE")
    d.rectangle((575, 155, 1025, 820), outline="#9FB3C8", width=2, fill="#F8FBFE")
    d.rectangle((1080, 155, 1530, 820), outline="#9FB3C8", width=2, fill="#F8FBFE")
    for x, label in [(225, "ESTUDIANTE"), (730, "SECRETARÍA"), (1235, "DIRECCIÓN DE CARRERA")]:
        d.rectangle((x - 150, 155, x + 150, 200), fill=NAVY)
        text_center(d, (x - 150, 155, x + 150, 200), label, FONT_14, WHITE)
    ellipse(d, (190, 245, 270, 325), "#DFF3F0", TEAL, text="Inicio", fnt=FONT_13)
    rounded(d, (140, 390, 340, 470), "#FFF7E1", GOLD, text="Registrar solicitud\ny requisitos", fnt=FONT_14)
    rounded(d, (640, 390, 920, 470), "#EAF2F8", NAVY, text="Validar identidad y\nregistro universitario", fnt=FONT_14)
    diamond = [(1235, 360), (1340, 430), (1235, 500), (1130, 430)]
    d.polygon(diamond, fill="#FDECEC", outline=RED)
    d.line(diamond + [diamond[0]], fill=RED, width=3)
    text_center(d, (1145, 390, 1325, 470), "¿Cumple\ncriterios?", FONT_13)
    rounded(d, (1140, 570, 1330, 650), "#EAF2F8", NAVY, text="Emitir resolución", fnt=FONT_14)
    rounded(d, (140, 670, 340, 750), "#DFF3F0", TEAL, text="Notificar resultado", fnt=FONT_14)
    ellipse(d, (190, 770, 270, 850), "#163A5F", NAVY, text="Fin", fnt=FONT_13)
    arrow(d, (230, 325), (230, 390)); arrow(d, (340, 430), (640, 430)); arrow(d, (920, 430), (1130, 430)); arrow(d, (1235, 500), (1235, 570), label="Sí"); arrow(d, (1140, 610), (340, 710)); arrow(d, (230, 750), (230, 770));
    arrow(d, (1340, 430), (1450, 430), label="No")
    rounded(d, (1350, 380, 1510, 480), "#FDECEC", RED, text="Solicitar\nsubsanación", fnt=FONT_13)
    arrow(d, (1430, 480), (920, 430), fill=RED, width=3)
    return save_img(img, "02_diagrama_actividad_casos_especiales.png")


def diag_sequence():
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title(d, "EASYDOC - Diagrama de secuencia: edición colaborativa de una política", "Persistencia versionada y sincronización de presencia mediante WebSocket")
    xs = [180, 470, 760, 1050, 1340]
    names = ["Editor A", "Angular", "FastAPI", "MongoDB", "Editor B"]
    colors = ["#EAF2F8", "#DFF3F0", "#FFF7E1", "#FDECEC", "#EAF2F8"]
    for x, name, color in zip(xs, names, colors):
        rounded(d, (x - 85, 150, x + 85, 215), color, NAVY, text=name, fnt=FONT_14)
        d.line((x, 215, x, 815), fill="#93A4B7", width=2)
    messages = [
        (260, 180, 470, "Editar nodo / conexión"),
        (325, 470, 760, "PUT /policies/{id} + versión"),
        (390, 760, 1050, "Guardar revisión y validar conflicto"),
        (455, 1050, 760, "201 / 409 + versión actual"),
        (520, 760, 470, "Confirmar guardado"),
        (590, 760, 1340, "WebSocket: operación y presencia"),
        (655, 1340, 470, "Aplicar operación remota"),
        (720, 470, 180, "Estado de sincronización"),
    ]
    for y, x1, x2, lab in messages:
        arrow(d, (x1, y), (x2, y), fill=TEAL if x1 < x2 else NAVY, width=3, label=lab)
    d.rounded_rectangle((620, 760, 900, 815), radius=12, outline=GOLD, width=3, fill="#FFF7E1")
    text_center(d, (630, 765, 890, 810), "Conflicto: se conserva la versión publicada y se informa al editor", FONT_12)
    return save_img(img, "03_diagrama_secuencia_colaboracion.png")


def diag_domain():
    img = Image.new("RGB", (1600, 940), "white")
    d = ImageDraw.Draw(img)
    title(d, "EASYDOC - Modelo lógico del dominio", "Entidades principales para usuarios, políticas, expedientes y trazabilidad")
    classes = [
        (85, 180, 350, 420, "Usuario", ["id", "nombre", "email", "roles[]", "activo"]),
        (470, 180, 760, 440, "Política", ["id", "nombre", "estado", "version", "diagrama", "formulario"]),
        (880, 180, 1150, 420, "Versión de política", ["id", "policy_id", "número", "contenido", "creado_en"]),
        (1250, 180, 1510, 420, "Flujo / tarea", ["id", "policy_id", "estado", "nodo_actual", "prioridad"]),
        (280, 610, 560, 850, "Documento", ["id", "título", "folder_id", "estado", "contenido"]),
        (720, 610, 1000, 850, "Versión documento", ["id", "document_id", "número", "autor", "creado_en"]),
        (1160, 610, 1460, 850, "Auditoría", ["id", "usuario_id", "acción", "recurso", "fecha"]),
    ]
    for x1, y1, x2, y2, name, attrs in classes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=14, outline=NAVY, width=3, fill="#F8FBFE")
        d.rectangle((x1, y1, x2, y1 + 46), fill=NAVY)
        text_center(d, (x1, y1, x2, y1 + 46), name, FONT_14, WHITE)
        y = y1 + 62
        for attr in attrs:
            d.text((x1 + 18, y), f"• {attr}", font=FONT_13, fill=INK)
            y += 29
    edges = [((350, 300), (470, 300), "crea"), ((760, 300), (880, 300), "versiona"), ((1150, 330), (1250, 330), "origina"), ((350, 420), (420, 610), "adjunta"), ((560, 730), (720, 730), "versiona"), ((560, 800), (1160, 800), "registra"), ((1510, 390), (1300, 610), "audita")]
    for st, en, lab in edges:
        arrow(d, st, en, fill=TEAL, width=3, label=lab)
    return save_img(img, "04_diagrama_dominio_easydoc.png")


def diag_components():
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title(d, "EASYDOC - Arquitectura de componentes", "Vista C4 de contenedores y servicios de la solución")
    groups = [(70, 160, 440, 790, "Clientes"), (555, 160, 1045, 790, "Servicios de aplicación"), (1160, 160, 1530, 790, "Datos e integración")]
    for x1, y1, x2, y2, name in groups:
        d.rounded_rectangle((x1, y1, x2, y2), radius=18, outline="#A6B9CC", width=3, fill="#F8FBFE")
        d.text((x1 + 20, y1 + 18), name, font=FONT_18, fill=NAVY)
    nodes = [
        (130, 280, 380, 365, "Web Angular 22\nMesa de control"), (130, 480, 380, 565, "Móvil Flutter\nConsulta y alertas"),
        (625, 245, 965, 330, "API FastAPI\nREST + JWT + WebSocket"), (625, 405, 965, 490, "Motor de políticas\nFormularios y workflow"), (625, 565, 965, 650, "Workers Celery\nDocumentos y reportes"),
        (1215, 230, 1470, 310, "MongoDB\nDatos y versiones"), (1215, 385, 1470, 465, "Redis\nBroker y caché"), (1215, 540, 1470, 620, "MinIO\nObjetos y archivos"), (1215, 680, 1470, 760, "IA local / LLM\nRiesgos y asistencia"),
    ]
    for x1, y1, x2, y2, lab in nodes:
        rounded(d, (x1, y1, x2, y2), "#DFF3F0" if x1 > 500 else "#EAF2F8", TEAL if x1 > 500 else NAVY, text=lab, fnt=FONT_14)
    for st, en, lab in [((380, 322), (625, 287), "HTTPS/WSS"), ((380, 522), (625, 445), "HTTPS"), ((965, 287), (1215, 270), "CRUD"), ((965, 445), (1215, 425), "pub/sub"), ((965, 605), (1215, 580), "S3 API"), ((965, 605), (1215, 720), "análisis")]:
        arrow(d, st, en, fill=GOLD, width=3, label=lab)
    return save_img(img, "05_diagrama_componentes_easydoc.png")


def diag_deployment():
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title(d, "EASYDOC - Vista de despliegue", "Entorno Docker para operación local y proyección de producción")
    d.rounded_rectangle((85, 180, 460, 760), radius=20, outline=NAVY, width=4, fill="#F8FBFE")
    d.text((120, 205), "Usuarios institucionales", font=FONT_18, fill=NAVY)
    rounded(d, (145, 300, 400, 385), "#EAF2F8", NAVY, text="Navegador web\nAngular", fnt=FONT_14)
    rounded(d, (145, 490, 400, 575), "#EAF2F8", NAVY, text="Dispositivo móvil\nFlutter", fnt=FONT_14)
    d.rounded_rectangle((575, 150, 1500, 800), radius=20, outline=TEAL, width=4, fill="#FCFFFE")
    d.text((610, 175), "Red Docker / infraestructura de aplicación", font=FONT_18, fill=TEAL)
    nodes = [
        (640, 290, 890, 370, "FastAPI\ncontenedor backend"), (640, 500, 890, 580, "Celery worker\nprocesos asíncronos"),
        (1040, 240, 1300, 320, "MongoDB\ncolecciones e índices"), (1040, 390, 1300, 470, "Redis\nbroker y caché"), (1040, 540, 1300, 620, "MinIO\nobjetos persistentes"), (1040, 690, 1300, 770, "ONLYOFFICE / IA\nservicios integrados"),
    ]
    for x1, y1, x2, y2, lab in nodes:
        rounded(d, (x1, y1, x2, y2), "#DFF3F0", TEAL, text=lab, fnt=FONT_14)
    for st, en, lab in [((400, 342), (640, 330), "HTTPS"), ((400, 532), (640, 540), "API"), ((890, 330), (1040, 280), "datos"), ((890, 330), (1040, 430), "eventos"), ((890, 540), (1040, 580), "archivos"), ((890, 540), (1040, 730), "tareas")]:
        arrow(d, st, en, fill=GOLD, width=3, label=lab)
    return save_img(img, "06_diagrama_despliegue_easydoc.png")


def make_drawio(diagram_names):
    mxfile = ET.Element("mxfile", {"host": "app.diagrams.net", "agent": "EASYDOC document generator", "pages": str(len(diagram_names))})
    for idx, name in enumerate(diagram_names, 1):
        diagram = ET.SubElement(mxfile, "diagram", {"name": name, "id": f"easydoc-{idx}"})
        model = ET.SubElement(diagram, "mxGraphModel", {"grid": "1", "page": "1", "pageWidth": "1600", "pageHeight": "900", "math": "0", "shadow": "0"})
        root = ET.SubElement(model, "root")
        ET.SubElement(root, "mxCell", {"id": "0"})
        ET.SubElement(root, "mxCell", {"id": "1", "parent": "0"})
        title_cell = ET.SubElement(root, "mxCell", {"id": f"{idx}-title", "value": name, "style": "rounded=1;whiteSpace=wrap;html=1;fillColor=#163A5F;fontColor=#FFFFFF;fontSize=18;fontStyle=1;", "vertex": "1", "parent": "1"})
        ET.SubElement(title_cell, "mxGeometry", {"x": "80", "y": "60", "width": "1440", "height": "60", "as": "geometry"})
        labels = ["Usuario / actor", "EASYDOC", "Servicio", "Datos / resultado"]
        for j, label in enumerate(labels, 1):
            cell = ET.SubElement(root, "mxCell", {"id": f"{idx}-{j}", "value": label, "style": "rounded=1;whiteSpace=wrap;html=1;fillColor=#DFF3F0;strokeColor=#0F766E;fontSize=14;", "vertex": "1", "parent": "1"})
            ET.SubElement(cell, "mxGeometry", {"x": str(140 + (j - 1) * 350), "y": "330", "width": "220", "height": "90", "as": "geometry"})
        for j in range(1, 4):
            edge = ET.SubElement(root, "mxCell", {"id": f"{idx}-e{j}", "value": "", "style": "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;", "edge": "1", "parent": "1", "source": f"{idx}-{j}", "target": f"{idx}-{j+1}"})
            ET.SubElement(edge, "mxGeometry", {"relative": "1", "as": "geometry"})
    tree = ET.ElementTree(mxfile)
    ET.indent(tree, space="  ")
    tree.write(DRAWIO, encoding="utf-8", xml_declaration=True)


def prepare_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths = [diag_use_cases(), diag_activity(), diag_sequence(), diag_domain(), diag_components(), diag_deployment()]
    make_drawio([
        "Casos de uso EASYDOC",
        "Actividad - casos especiales académicos",
        "Secuencia - colaboración de políticas",
        "Dominio documental",
        "Componentes EASYDOC",
        "Despliegue Docker",
    ])
    return paths


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    cell.width = Inches(width_inches)


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        item = borders.find(tag)
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "8AA0B4")


def set_font(paragraph, size=10.0, color=INK, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        run.bold = bold
        run.italic = italic


def add_body(doc, text, italic=False):
    p = doc.add_paragraph(style="normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Inches(0.28)
    pf.line_spacing = 1.35
    pf.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    r.italic = italic
    return p


def add_bullets(doc, values):
    for value in values:
        p = doc.add_paragraph(style="normal")
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "28")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        p_pr.append(num_pr)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(value)
        set_font(p, 11)


def add_heading(doc, text, level=1, page_before=False):
    if page_before:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def patch_toc_field(docx_path):
    """Use outline levels because the retained Word template is localized."""
    temp = docx_path.with_name(docx_path.stem + "_toc_patch.docx")
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    descriptions = [
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Elemento gráfico de la carátula original preservada del modelo.",
        "Diagrama general de casos de uso de EASYDOC.",
        "Flujo de actividad de la política Casos Especiales Académicos.",
        "Secuencia de colaboración, versionado y control de conflictos.",
        "Modelo lógico del dominio documental y de workflow.",
        "Componentes de EASYDOC y sus integraciones principales.",
        "Despliegue lógico de EASYDOC con servicios Docker.",
        "Mesa de Control Académica de EASYDOC.",
        "Biblioteca de políticas académicas publicadas.",
        "Política publicada en el diseñador UML de EASYDOC.",
        "Pantalla inicial y Mesa de Control Académica.",
        "Lienzo del diseñador UML antes de registrar una política.",
        "Gestión de carpetas, filtros y búsqueda de expedientes.",
        "Panel de riesgos y predicciones de EASYDOC.",
    ]
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for node in root.xpath(".//w:instrText", namespaces=ns):
                    if node.text and "TOC" in node.text:
                        node.text = ' TOC \\o "1-4" \\h \\z \\u '
                for index, node in enumerate(root.xpath(".//wp:docPr", namespaces=ns)):
                    description = descriptions[index] if index < len(descriptions) else "Recurso visual de EASYDOC."
                    node.set("descr", description)
                    node.set("title", description)
                for table in root.xpath(".//w:tbl", namespaces=ns):
                    first_row = table.find(qn("w:tr"))
                    if first_row is not None:
                        row_properties = first_row.find(qn("w:trPr"))
                        if row_properties is None:
                            row_properties = OxmlElement("w:trPr")
                            first_row.insert(0, row_properties)
                        if row_properties.find(qn("w:tblHeader")) is None:
                            header = OxmlElement("w:tblHeader")
                            header.set(qn("w:val"), "true")
                            row_properties.append(header)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "word/settings.xml":
                root = etree.fromstring(data)
                for node in root.xpath(".//w:updateFields", namespaces=ns):
                    node.set(qn("w:val"), "false")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    temp.replace(docx_path)


def add_caption(doc, text):
    p = doc.add_paragraph(style="normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    set_font(p, 10.5, italic=True)
    return p


def add_figure(doc, image_path, caption, width=6.25):
    p = doc.add_paragraph(style="normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if widths is None:
        widths = [6.35 / len(headers)] * len(headers)
    set_table_width(table, widths)
    set_table_borders(table)
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, 9.5, color=WHITE, bold=True)
    for record in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, record):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                set_font(p, 9.0)
    set_table_width(table, widths)
    after = doc.add_paragraph(style="normal")
    after.paragraph_format.space_after = Pt(4)
    return table


def add_cu_table(doc, code, name, purpose, actors, precondition, flow, result, exception):
    add_heading(doc, f"{code}. {name}", 4)
    rows = [
        ("Nombre de CU", f"{code}. {name}"),
        ("Propósito", purpose),
        ("Actores", actors),
        ("Actor iniciador", actors.split(",")[0]),
        ("Precondición", precondition),
        ("Flujo principal", flow),
        ("Postcondición", result),
        ("Excepción", exception),
    ]
    add_table(doc, ["Elemento", "Especificación"], rows, [1.35, 5.0])


def clean_template(doc):
    first_heading = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip() == "PERFIL")
    body = doc._element.body
    children = list(body)
    cut = children.index(first_heading._p)
    for child in children[cut:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    # The retained model stores some Word margins as decimal strings (for
    # example, 1440.0000000000002). Word accepts them, while python-docx needs
    # integral twips before it can create new tables. Rounding preserves 1-inch
    # margins and leaves the visual geometry unchanged.
    for sect_pr in doc._element.iter(qn("w:sectPr")):
        pg_mar = sect_pr.find(qn("w:pgMar"))
        if pg_mar is not None:
            for attr in ("top", "right", "bottom", "left", "header", "footer", "gutter"):
                key = qn(f"w:{attr}")
                value = pg_mar.get(key)
                if value and "." in value:
                    pg_mar.set(key, str(round(float(value))))
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build_document():
    prepare_assets()
    shutil.copy2(REFERENCE, FINAL)
    doc = Document(str(FINAL))
    clean_template(doc)
    doc.core_properties.title = "EASYDOC - Sistema de Gestión Documental Académica"
    doc.core_properties.subject = "Documento PUDS adaptado a EASYDOC"
    doc.core_properties.comments = "Contenido generado a partir de la estructura del documento modelo; carátula preservada."

    # PERFIL
    add_heading(doc, "PERFIL", 1)
    add_heading(doc, "INTRODUCCIÓN", 2)
    for t in [
        "La Dirección de Carrera requiere atender solicitudes académicas, conservar expedientes, aplicar políticas institucionales y dar seguimiento oportuno a cada trámite. Cuando estas tareas se resuelven mediante correos, carpetas aisladas o formularios no estandarizados, se incrementan las demoras, los documentos se duplican y resulta difícil comprobar quién realizó cada acción.",
        "EASYDOC es una plataforma web y móvil orientada a una Dirección de Carrera universitaria. Integra una mesa de control académica, biblioteca de políticas, diseñador de actividades UML, formularios dinámicos, repositorio documental, colaboración, auditoría, reportes y análisis de riesgos. Su propósito es convertir el flujo académico en un proceso trazable, gobernado y comprensible para estudiantes, secretarías, responsables y dirección.",
        "La solución aplica el Proceso Unificado de Desarrollo de Software (PUDS): está dirigida por casos de uso, centrada en la arquitectura y organizada por iteraciones. La descripción funcional se expresa con UML 2.5.1 para conservar una relación verificable entre los requisitos, el diseño, la implementación y las pruebas (Object Management Group [OMG], 2017; Jacobson et al., 1999).",
    ]:
        add_body(doc, t)
    add_heading(doc, "OBJETIVO GENERAL", 2)
    add_body(doc, "Implementar una plataforma integral de gestión documental académica que permita modelar y publicar políticas de negocio, registrar y seguir trámites, administrar expedientes y formularios, colaborar sobre los procesos, analizar riesgos operativos y producir evidencia auditable para la Dirección de Carrera.")
    add_heading(doc, "OBJETIVOS ESPECÍFICOS", 2)
    add_bullets(doc, [
        "Estandarizar las políticas académicas mediante diagramas de actividad UML y formularios asociados a requisitos documentales.",
        "Centralizar expedientes, versiones y permisos para asegurar disponibilidad, integridad y rastreabilidad de los documentos.",
        "Ofrecer una mesa de control que priorice solicitudes, muestre riesgos y permita a los responsables tomar decisiones oportunas.",
        "Aplicar autenticación JWT, control de acceso por roles y registros de auditoría para proteger información académica.",
        "Incorporar un asistente y analítica local que sugieran rutas, alerten cuellos de botella y faciliten reportes de gestión.",
    ])
    add_heading(doc, "DESCRIPCIÓN DEL PROBLEMA", 2)
    for t in [
        "Las solicitudes de certificados, homologaciones, becas, casos especiales y otros trámites académicos involucran documentos, validaciones, decisiones y plazos. En un flujo no integrado, el estudiante desconoce el estado real de su solicitud; las secretarías repiten validaciones; y Dirección no cuenta con una visión consolidada de la carga ni de los expedientes con riesgo de retraso.",
        "El problema se amplifica cuando una política cambia o un expediente debe ser corregido. Sin una versión formal del flujo, los participantes pueden aplicar reglas distintas. Sin un repositorio con historial, la documentación no es verificable. Sin trazabilidad y alertas, la intervención llega cuando el plazo ya está comprometido.",
        "EASYDOC resuelve esta situación mediante una única fuente de verdad para políticas, instancias de workflow, formularios y documentos. La plataforma separa la definición de una política de su ejecución, conserva revisiones, limita las acciones según el rol y emplea eventos para actualizar la interfaz y el historial institucional.",
    ]:
        add_body(doc, t)
    add_heading(doc, "ALCANCE", 2)
    modules = [
        ("Acceso y gobierno", "Autenticación, roles, administración de usuarios, configuración y auditoría de acciones."),
        ("Políticas y UML", "Biblioteca de políticas, diseñador de actividades, publicación, versionado y exportación de diagramas."),
        ("Formularios y workflow", "Campos configurables, requisitos adjuntos, inicio y avance de trámites según la política publicada."),
        ("Expedientes", "Carpetas, documentos, versiones, previsualización, permisos y colaboración."),
        ("Control y riesgos", "Mesa de control, métricas, alertas de cuello de botella y recomendaciones de IA local."),
        ("Asistencia y reportes", "Asistente EASY, búsqueda contextual y generación de reportes de gestión."),
        ("Canal móvil", "Acceso móvil a información priorizada, sesión segura y continuidad de consulta."),
    ]
    add_table(doc, ["Módulo", "Alcance funcional"], modules, [1.65, 4.7])
    add_figure(doc, ASSETS / "01_diagrama_casos_uso_easydoc.png", "Figura 1. Diagrama general de casos de uso de EASYDOC.")

    # MARCO TEÓRICO
    add_heading(doc, "MARCO TEÓRICO", 1, page_before=True)
    sections = [
        ("PUDS", "El Proceso Unificado organiza el trabajo en inicio, elaboración, construcción y transición. Su valor para EASYDOC radica en que los casos de uso académicos orientan las iteraciones, la arquitectura reduce riesgos de integración y las pruebas validan cada incremento antes de su liberación (Jacobson et al., 1999)."),
        ("UML y modelado de políticas", "UML proporciona una notación estandarizada para visualizar, especificar y documentar artefactos de software. En EASYDOC los diagramas de actividad describen los pasos, decisiones y responsables de una política académica; los diagramas de clases y componentes mantienen la coherencia con el diseño técnico (OMG, 2017)."),
        ("Gestión documental y trazabilidad", "Un sistema de gestión documental debe administrar el ciclo de vida del expediente: registro, clasificación, acceso, modificación, versiones, conservación y consulta. En EASYDOC los documentos se relacionan con carpetas, versiones y registros de auditoría para permitir que una decisión académica pueda reconstruirse a partir de su evidencia."),
        ("Colaboración en tiempo real", "Los canales WebSocket permiten comunicación bidireccional entre cliente y servidor. El uso de salas por política evita mezclar operaciones de procesos distintos y permite informar presencia, cambios y conflictos de revisión sin recargar la interfaz."),
        ("Inteligencia artificial local", "La analítica de EASYDOC emplea datos de flujos institucionales para indicar carga, anomalías documentales y rutas recomendadas. La recomendación no sustituye al responsable académico: expone evidencia, riesgo y una acción sugerida que debe ser validada por el rol competente."),
        ("Arquitectura de servicios", "Angular facilita una interfaz web modular; FastAPI expone contratos REST y WebSocket con validación basada en tipos; MongoDB almacena estructuras documentales flexibles; Redis sirve como caché y broker; y MinIO ofrece almacenamiento compatible con S3. Esta combinación prioriza modularidad, mantenibilidad e integración por interfaces (Angular, 2026; FastAPI, 2026; MongoDB, 2026)."),
    ]
    for idx, (head, body) in enumerate(sections, 1):
        add_heading(doc, f"2.{idx}. {head}", 2)
        add_body(doc, body)
    add_heading(doc, "2.7. Comparativa de tecnologías", 2)
    add_table(doc, ["Necesidad", "Tecnología seleccionada", "Criterio de selección"], [
        ("Cliente web", "Angular 22", "Componentes, rutas, formularios reactivos y ecosistema institucional."),
        ("API", "FastAPI + Pydantic", "Contratos tipados, documentación OpenAPI y soporte asíncrono."),
        ("Persistencia", "MongoDB + Beanie", "Modelos documentales y evolución de formularios sin migraciones rígidas."),
        ("Eventos", "WebSocket + Redis", "Actualización colaborativa y separación de tareas en segundo plano."),
        ("Archivos", "MinIO / S3", "Objetos versionables y almacenamiento desacoplado."),
        ("Móvil", "Flutter", "Base multiplataforma para consulta y continuidad operativa."),
    ], [1.3, 1.6, 3.45])

    # REQUISITOS
    add_heading(doc, "CAPÍTULO 2 FLUJO DE TRABAJO: CAPTURA DE REQUISITOS", 1, page_before=True)
    add_heading(doc, "Identificar actores y casos de uso", 2)
    add_body(doc, "Los actores fueron identificados desde las responsabilidades reales de una Dirección de Carrera. El acceso a cada capacidad depende del rol, de la política publicada y de los permisos del expediente asociado.")
    add_table(doc, ["Actor", "Responsabilidad", "Interacción principal"], [
        ("Director / administrador", "Gobierna usuarios, políticas, publicación, reportes y control institucional.", "Biblioteca de políticas, gobierno, riesgos y auditoría."),
        ("Secretaría / funcionario", "Valida requisitos, ejecuta tareas asignadas y administra expedientes.", "Mesa de control, formularios y repositorio."),
        ("Estudiante", "Registra solicitudes, adjunta requisitos y consulta resultados.", "Trámites, expedientes y notificaciones."),
        ("Servicio IA local", "Analiza patrones y propone una ruta o alerta contextual.", "Riesgos, asistente y reportes."),
    ], [1.3, 2.75, 2.3])
    add_heading(doc, "Priorizar casos de uso", 2)
    add_table(doc, ["Ciclo", "Casos priorizados", "Criterio"], [
        ("Ciclo 1", "CU01 a CU09", "Operación mínima: identidad, políticas, formularios, workflow, expediente y trazabilidad."),
        ("Ciclo 2", "CU10 a CU14", "Optimización: IA local, reportes, gobierno ampliado y canal móvil."),
    ], [1.1, 1.7, 3.55])
    add_heading(doc, "Especificar casos de uso", 2)
    cu_specs = [
        ("CU01", "Gestionar acceso, usuarios y roles", "Administrar identidades, roles y sesión segura.", "Director / administrador", "Cuenta autorizada.", "El administrador registra o modifica un usuario, asigna rol y el sistema valida la operación antes de persistirla.", "La cuenta queda disponible con permisos consistentes.", "Correo duplicado, rol no permitido o sesión vencida."),
        ("CU02", "Administrar políticas académicas", "Crear, consultar, publicar o archivar una política.", "Director / administrador", "Sesión con permiso de gestión.", "El responsable define nombre, estado y metadatos; la biblioteca conserva la política y su versión vigente.", "Política creada, actualizada o publicada.", "Nombre inválido, permiso insuficiente o conflicto de versión."),
        ("CU03", "Diseñar diagrama de actividades UML", "Representar la secuencia, decisiones y responsables de una política.", "Director / administrador", "Política en edición.", "El usuario agrega inicio, actividades, decisiones, fork, join y fin; enlaza los nodos y registra el diagrama.", "Diagrama válido asociado a la política.", "Nodos desconectados, ausencia de inicio/fin o enlace inválido."),
        ("CU04", "Configurar formularios y requisitos", "Definir campos y adjuntos exigidos por la política.", "Director / administrador", "Política y nodo de tarea disponibles.", "El usuario agrega campos, opciones, obligatoriedad y archivos requeridos; la plataforma valida la definición.", "Formulario reutilizable publicado.", "Campo sin etiqueta, requisito duplicado o tipo no admitido."),
        ("CU05", "Registrar y avanzar un trámite", "Ejecutar una solicitud conforme a la política publicada.", "Estudiante, Secretaría", "Política publicada y requisitos iniciales disponibles.", "Se inicia la instancia, se captura el formulario y el motor dirige la tarea al siguiente nodo según la regla definida.", "Trámite actualizado con estado y nodo actual.", "Requisitos incompletos, política no publicada o decisión sin ruta."),
        ("CU06", "Gestionar expedientes documentales", "Centralizar documentos, carpetas, metadatos y permisos.", "Secretaría, Director", "Usuario autenticado.", "El usuario carga, clasifica, busca, consulta o versiona un documento; se registra el evento correspondiente.", "Expediente actualizado y trazable.", "Archivo inválido, permiso insuficiente o almacenamiento no disponible."),
        ("CU07", "Editar documentos colaborativamente", "Permitir trabajo concurrente controlado en un recurso compartido.", "Secretaría, Director", "Documento editable y permiso de modificación.", "Los participantes abren el documento, realizan cambios y la plataforma conserva una nueva versión y actividad de colaboración.", "Versión y participación registradas.", "Conflicto de edición, conexión perdida o acceso de solo lectura."),
        ("CU08", "Consultar mesa de control", "Visualizar solicitudes, estados, carga y prioridades.", "Director, Secretaría", "Sesión autenticada.", "El actor abre la mesa de control, filtra información y consulta las alertas o cola de documentos recientes.", "Visión consolidada de la operación.", "Sin datos para el filtro o conectividad temporalmente no disponible."),
        ("CU09", "Gestionar notificaciones y trazabilidad", "Informar eventos y conservar evidencia de acciones.", "Todos los roles", "Evento de negocio generado.", "El sistema registra el evento, notifica al destinatario por el canal disponible y permite consultar la auditoría.", "Notificación y registro auditables.", "Destinatario no disponible; se conserva para reintento o consulta."),
        ("CU10", "Analizar riesgos operativos con IA local", "Detectar anomalías, rutas lentas y oportunidades de reasignación.", "Director", "Datos históricos suficientes.", "El usuario selecciona periodo y trámite; el servicio calcula métricas, muestra riesgo y explica la recomendación.", "Riesgo priorizado con acción sugerida.", "Datos insuficientes o modelo local no disponible."),
        ("CU11", "Interactuar con Asistente EASY", "Resolver consultas sobre políticas, expedientes y rutas de atención.", "Estudiante, Secretaría, Director", "Sesión o contexto permitido.", "El actor formula una consulta y el asistente devuelve una respuesta contextual, ruta propuesta o pregunta de aclaración.", "Interacción registrada sin alterar la política de forma automática.", "Consulta ambigua, contexto no autorizado o sin resultados."),
        ("CU12", "Generar reportes de gestión", "Obtener reportes operativos y exportables.", "Director", "Permiso de reporte.", "El usuario define periodo y criterio; el sistema consolida métricas, produce una vista y habilita la exportación.", "Reporte disponible para consulta.", "Filtro inválido, datos insuficientes o formato no soportado."),
        ("CU13", "Administrar gobierno y configuración", "Gestionar parámetros, perfiles y políticas de seguridad.", "Director / administrador", "Rol administrador.", "El administrador revisa la configuración, ajusta parámetros permitidos y consulta controles de cumplimiento.", "Configuración institucional actualizada.", "Cambio no permitido o autorización insuficiente."),
        ("CU14", "Operar desde el canal móvil", "Consultar información priorizada y mantener sesión segura desde móvil.", "Director, Secretaría", "Aplicación instalada y usuario autorizado.", "El usuario inicia sesión, consulta tablero y recibe información de tareas o alertas relevantes.", "Consulta móvil coherente con el perfil del usuario.", "Dispositivo no autorizado, biometría fallida o sesión expirada."),
    ]
    add_heading(doc, "Ciclo #1", 3)
    for spec in cu_specs[:9]: add_cu_table(doc, *spec)
    add_heading(doc, "Ciclo #2", 3)
    for spec in cu_specs[9:]: add_cu_table(doc, *spec)
    add_heading(doc, "Estructurar el modelo de casos de uso", 2)
    add_body(doc, "Los casos de uso se agrupan por responsabilidad y se conectan a los módulos de EasyDoc. Esta estructura permite que cada caso pueda rastrearse a clases, componentes, rutas de interfaz y pruebas funcionales.")
    add_figure(doc, ASSETS / "02_diagrama_actividad_casos_especiales.png", "Figura 2. Flujo de actividad de la política Casos Especiales Académicos.")

    # ANÁLISIS
    add_heading(doc, "CAPÍTULO 3: ANÁLISIS", 1, page_before=True)
    add_heading(doc, "Análisis de arquitectura", 2)
    add_body(doc, "La arquitectura separa presentación, aplicación, dominio e infraestructura. Las rutas de FastAPI reciben solicitudes REST y WebSocket; los servicios aplican reglas; los modelos y esquemas validan datos; y los adaptadores acceden a MongoDB, Redis, MinIO y servicios externos. Esta separación disminuye el acoplamiento y hace que los módulos puedan evolucionar sin alterar la interfaz pública.")
    add_heading(doc, "Identificar paquetes", 3)
    add_table(doc, ["Paquete", "Responsabilidad", "Relación con CU"], [
        ("Frontend Angular", "Vistas, guardas, componentes y servicios de API.", "CU01-CU14"),
        ("Routers FastAPI", "Contratos REST/WebSocket, autenticación y validación de solicitudes.", "CU01-CU13"),
        ("Servicios", "Reglas de políticas, workflow, archivos, auditoría, IA y reportes.", "CU02-CU12"),
        ("Modelos y esquemas", "Persistencia documental y contratos Pydantic.", "CU01-CU13"),
        ("Workers", "Procesamiento diferido de reportes, notificaciones y tareas de análisis.", "CU06, CU09, CU10, CU12"),
    ], [1.35, 2.85, 1.8])
    add_heading(doc, "Análisis de casos de uso", 2)
    add_heading(doc, "Colaboración y persistencia versionada", 3)
    add_body(doc, "La edición de una política combina una actualización REST versionada con un canal de colaboración. El guardado incluye la versión esperada; si otra persona publicó una revisión antes, se devuelve un conflicto que evita sobrescribir cambios. Las operaciones y la presencia se difunden solo a las sesiones de esa política.")
    add_figure(doc, ASSETS / "03_diagrama_secuencia_colaboracion.png", "Figura 3. Secuencia de colaboración, versionado y control de conflictos.")
    add_heading(doc, "Análisis de clases", 2)
    add_body(doc, "El dominio se construye alrededor de la identidad, la política y el expediente. Una política se revisa como unidad versionada; una instancia de flujo conserva su estado; y un documento obtiene versiones propias. La auditoría conecta cada recurso con la acción y el usuario que la ejecutó.")
    add_figure(doc, ASSETS / "04_diagrama_dominio_easydoc.png", "Figura 4. Modelo lógico del dominio documental y de workflow.")
    add_heading(doc, "Análisis de riesgos", 2)
    add_table(doc, ["Riesgo", "Impacto", "Mitigación implementada"], [
        ("Edición simultánea", "Pérdida de cambios en políticas.", "Versionado optimista, respuesta 409 y difusión por sala WebSocket."),
        ("Documento incompleto", "Retraso o decisión no sustentada.", "Formularios con obligatoriedad y requisitos de archivo independientes."),
        ("Acceso no autorizado", "Exposición de información académica.", "JWT, roles, guardas de rutas y registro de auditoría."),
        ("Cuello de botella", "Incumplimiento de plazos de atención.", "Mesa de control, indicadores y recomendación local de riesgo."),
        ("Falla de infraestructura", "Interrupción de la operación.", "Servicios desacoplados en contenedores y datos persistentes."),
    ], [1.45, 2.2, 2.35])

    # DISEÑO
    add_heading(doc, "CAPÍTULO 4 FLUJO DE TRABAJO: DISEÑO", 1, page_before=True)
    add_heading(doc, "Diseño de arquitectura C4", 2)
    add_heading(doc, "Nivel 1: contexto institucional", 3)
    add_body(doc, "EASYDOC se sitúa entre los miembros de la Dirección de Carrera y sus expedientes académicos. Los usuarios consultan, registran y gobiernan procesos desde la web o el móvil; los servicios de aplicación realizan validaciones, persistencia, procesamiento y análisis.")
    add_heading(doc, "Nivel 2: contenedores y componentes", 3)
    add_figure(doc, ASSETS / "05_diagrama_componentes_easydoc.png", "Figura 5. Componentes de EASYDOC y sus integraciones principales.")
    add_heading(doc, "Diseño de datos", 2)
    add_heading(doc, "Modelo de datos lógico", 3)
    add_table(doc, ["Colección / entidad", "Propósito", "Relaciones relevantes"], [
        ("users y roles", "Identidad, permisos y perfiles institucionales.", "Usuario posee roles y genera auditoría."),
        ("policies y policy_versions", "Políticas UML, formularios, estado y revisiones.", "Política posee historial y origina workflows."),
        ("workflows y workflow_tasks", "Instancias de trámite, nodo actual y tareas.", "Workflow referencia política y expedientes."),
        ("documents, folders y document_versions", "Archivos, clasificación y control de versiones.", "Documento pertenece a carpeta y posee versiones."),
        ("audit_logs y notifications", "Trazabilidad y comunicación de eventos.", "Referencian usuario, recurso o trámite."),
        ("conversations y collaboration_sessions", "Contexto de asistencia y presencia colaborativa.", "Se vinculan al usuario y a la política/documento."),
    ], [1.7, 2.4, 1.9])
    add_heading(doc, "Diseño físico y seguridad", 3)
    add_body(doc, "MongoDB conserva los documentos de negocio y sus índices; Redis provee caché y cola de trabajo; MinIO concentra archivos binarios; y los workers realizan tareas que no deben bloquear la respuesta web. Las credenciales se gestionan mediante variables de entorno y las operaciones sensibles quedan sujetas a autorización y auditoría.")
    add_heading(doc, "Vista de despliegue", 2)
    add_figure(doc, ASSETS / "06_diagrama_despliegue_easydoc.png", "Figura 6. Despliegue lógico de EASYDOC con servicios Docker.")
    add_heading(doc, "Diseño de experiencia de usuario", 2)
    add_body(doc, "La interfaz emplea una jerarquía institucional: navegación lateral para los módulos, barra superior de búsqueda y estado institucional, tarjetas de control y paneles especializados. La biblioteca y el diseñador reducen la complejidad de las políticas al mostrar el flujo, las propiedades y los requisitos en una misma pantalla.")
    add_figure(doc, ASSETS / "captura_01_mesa_control.png", "Figura 7. Mesa de Control Académica de EASYDOC.", width=6.25)

    # IMPLEMENTACIÓN
    add_heading(doc, "CAPÍTULO 5 FLUJO DE TRABAJO: IMPLEMENTACIÓN", 1, page_before=True)
    add_heading(doc, "Elección de plataforma de desarrollo de software", 2)
    add_heading(doc, "Lenguajes y frameworks", 3)
    add_body(doc, "El frontend se implementa con Angular 22 y TypeScript; el backend se desarrolla con Python, FastAPI y Pydantic; y el canal móvil utiliza Flutter y Dart. La selección responde a la necesidad de interfaces modulares, contratos explícitos, validación de datos y una experiencia multiplataforma.")
    add_heading(doc, "Base de datos e infraestructura", 3)
    add_body(doc, "MongoDB almacena recursos con estructura variable, como diagramas y formularios; Redis actúa como broker y caché; MinIO aporta almacenamiento de objetos; y Celery procesa tareas asíncronas. Docker Compose describe el conjunto de servicios y permite ejecutar el entorno local de manera reproducible (Docker, 2026; MongoDB, 2026).")
    add_heading(doc, "Componentes implementados", 3)
    add_table(doc, ["Capa", "Componentes implementados", "Resultado"], [
        ("Web", "Dashboard, repositorio, editor, políticas, UML, asistente, riesgos, reportes, auditoría y gobierno.", "Navegación institucional y vistas especializadas."),
        ("Backend", "Routers, servicios, modelos, esquemas, autenticación, auditoría, reportes y colaboración.", "API REST/WebSocket con reglas de dominio."),
        ("Datos", "Modelos Beanie, seeds, índices y almacenamiento de objetos.", "Persistencia y demostración con datos académicos."),
        ("Móvil", "Sesión, biometría, tablero y navegación base.", "Acceso portátil y seguro a información priorizada."),
    ], [1.1, 3.5, 1.4])
    add_heading(doc, "Rutas y servicios principales", 3)
    add_table(doc, ["Interfaz", "Servicio", "Finalidad"], [
        ("/policies", "POST/GET/PUT /api/policies", "Crear, consultar y versionar políticas."),
        ("/uml-designer", "WebSocket de colaboración de políticas", "Difundir presencia y operaciones por política."),
        ("/repository", "API de documentos y almacenamiento S3", "Gestionar archivos, versiones y permisos."),
        ("/predictions", "Servicio de predicción local", "Exponer riesgos, anomalías y rutas sugeridas."),
        ("/reports", "Generador de reportes", "Consolidar y exportar información de gestión."),
        ("/audit", "Servicio de auditoría", "Consultar historial verificable de acciones."),
    ], [1.35, 2.6, 2.05])
    add_figure(doc, ASSETS / "captura_02_biblioteca_politicas.png", "Figura 8. Biblioteca de políticas académicas publicadas.", width=6.25)
    add_figure(doc, ASSETS / "captura_04_politica_publicada.png", "Figura 9. Política publicada en el diseñador UML de EASYDOC.", width=6.25)

    # TESTING AND MANUAL
    add_heading(doc, "CAPÍTULO 6 FLUJO DE TRABAJO: PRUEBAS", 1, page_before=True)
    add_heading(doc, "Estrategia de pruebas", 2)
    add_body(doc, "La validación se organiza por casos de uso, contratos de API, reglas de dominio y recorridos de interfaz. El repositorio contiene pruebas para autenticación, usuarios, roles, documentos, colaboración, políticas, predicciones, reportes, notificaciones, auditoría y modelos de datos. La matriz siguiente consolida los recorridos funcionales representativos de la versión entregada.")
    add_table(doc, ["ID", "Escenario", "Resultado esperado", "Estado"], [
        ("P-01", "Acceso con rol autorizado", "Se crea sesión y se habilitan únicamente las rutas del perfil.", "Conforme"),
        ("P-02", "Registro y guardado de política", "La política aparece en la biblioteca con su versión.", "Conforme"),
        ("P-03", "Edición de diagrama UML", "Nodos, enlaces y etiquetas se conservan al recargar.", "Conforme"),
        ("P-04", "Formulario con requisito obligatorio", "No se avanza mientras falte evidencia requerida.", "Conforme"),
        ("P-05", "Inicio y avance de trámite", "El motor actualiza nodo, estado, prioridad e historial.", "Conforme"),
        ("P-06", "Carga y versionado documental", "El expediente conserva archivo, metadatos y nueva versión.", "Conforme"),
        ("P-07", "Colaboración de política", "La operación se difunde a la sala y el conflicto se informa.", "Conforme"),
        ("P-08", "Consulta de mesa de control", "Se muestran solicitudes, documentos recientes, alertas y prioridades.", "Conforme"),
        ("P-09", "Predicción de riesgo", "Se visualizan anomalías, ruta sugerida e insight de reasignación.", "Conforme"),
        ("P-10", "Reporte de gestión", "Se consolida la información solicitada y se habilita exportación.", "Conforme"),
    ], [0.65, 2.0, 2.9, 0.8])
    test_sections = [
        ("Pruebas de autenticación y acceso", "Se comprobó la navegación protegida por guardas y la diferenciación de vistas según rol. Una sesión segura permite acceder a módulos autorizados y bloquea las rutas administrativas cuando el perfil no corresponde."),
        ("Pruebas de gestión de políticas y formularios", "Se verificó el registro de políticas, la creación de diagramas, la persistencia de formularios, la publicación y la recuperación desde la biblioteca. La interfaz comunica claramente el estado de borrador o edición guardada."),
        ("Pruebas de workflow y expedientes", "Se validó que una solicitud se inicie únicamente con una política publicada, que los requisitos obligatorios se controlen y que los cambios en documentos originen versiones y auditoría."),
        ("Pruebas de colaboración y trazabilidad", "El canal de política autentica la sesión, agrupa usuarios por recurso y expone cambios de presencia. El control de versión evita reemplazar una revisión más reciente de otro editor."),
        ("Pruebas de riesgos, asistente y reportes", "Las vistas de IA local presentan indicadores de tiempo, anomalías, rutas sugeridas y recomendaciones. La recomendación se muestra como apoyo a la decisión y no modifica un trámite sin intervención humana."),
    ]
    for i, (name, content) in enumerate(test_sections, 1):
        add_heading(doc, f"6.{i}. {name}", 2)
        add_body(doc, content)
    add_heading(doc, "Manual de usuario de EASYDOC", 2)
    add_heading(doc, "Acceder y reconocer la mesa de control", 3)
    add_body(doc, "Después de iniciar sesión, el usuario encuentra la Mesa de Control Académica. Las tarjetas superiores resumen solicitudes, formularios, archivos, riesgos y el consejo de IA. Use la barra lateral para abrir el módulo requerido y la búsqueda superior para localizar un estudiante, expediente, trámite o política.")
    add_figure(doc, ASSETS / "captura_01_mesa_control.png", "Figura 10. Pantalla inicial y Mesa de Control Académica.", width=6.25)
    add_heading(doc, "Crear o editar una política", 3)
    add_body(doc, "Abra Políticas de negocio y seleccione Nueva política o Editar. En el Diseñador de actividades UML escriba el nombre del proceso, arrastre Inicio, Actividad, Decisión, Fork, Join y Fin al lienzo, conecte los elementos y configure las propiedades. Registre la política cuando el flujo sea válido; podrá volver a abrirla desde la biblioteca.")
    add_figure(doc, ASSETS / "captura_03_disenador_uml.png", "Figura 11. Lienzo del diseñador UML antes de registrar una política.", width=6.25)
    add_heading(doc, "Administrar expedientes y documentos", 3)
    add_body(doc, "Abra Expedientes para elegir una carpeta, filtrar por tipo, autor o etiqueta y cargar archivos. Utilice la búsqueda semántica cuando necesite recuperar contenido por contexto. Antes de modificar un recurso, verifique que su rol tenga permiso de edición; cada operación quedará disponible para trazabilidad.")
    add_figure(doc, ASSETS / "captura_05_expedientes.png", "Figura 12. Gestión de carpetas, filtros y búsqueda de expedientes.", width=6.25)
    add_heading(doc, "Consultar riesgos y tomar decisiones", 3)
    add_body(doc, "Abra Riesgos IA, seleccione el periodo y el tipo de trámite. La pantalla presenta el tiempo promedio, las anomalías documentales, el ahorro proyectado, la ruta recomendada y los cuellos de botella. Revise el insight y confirme con el área responsable antes de reasignar recursos o ajustar una política.")
    add_figure(doc, ASSETS / "captura_06_riesgos_ia.png", "Figura 13. Panel de riesgos y predicciones de EASYDOC.", width=6.25)
    add_heading(doc, "Cierre de la operación", 3)
    add_body(doc, "Al finalizar una tarea, compruebe que el estado del trámite y los documentos se hayan actualizado. Si un usuario informa una inconsistencia, consulte Trazabilidad para identificar el evento, el recurso y el responsable. Las políticas se deben publicar únicamente después de validar su diagrama, formulario y responsables académicos.")

    # REFERENCES
    add_heading(doc, "BIBLIOGRAFÍAS", 1, page_before=True)
    references = [
        "Angular. (2026). What is Angular? https://angular.dev/overview",
        "Celery. (2026). Celery documentation. https://docs.celeryq.dev/",
        "Docker. (2026). Docker documentation. https://docs.docker.com/",
        "FastAPI. (2026). FastAPI documentation. https://fastapi.tiangolo.com/",
        "Fielding, R. T. (2000). Architectural styles and the design of network-based software architectures [Doctoral dissertation, University of California, Irvine].",
        "Flutter. (2026). Flutter documentation. https://docs.flutter.dev/",
        "Jacobson, I., Booch, G., & Rumbaugh, J. (1999). The unified software development process. Addison-Wesley.",
        "MinIO. (2026). MinIO object storage documentation. https://min.io/docs/",
        "MongoDB. (2026). MongoDB documentation. https://www.mongodb.com/docs/",
        "Object Management Group. (2017). Unified Modeling Language (UML), version 2.5.1. https://www.omg.org/spec/UML/2.5.1/",
        "Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner’s approach (9th ed.). McGraw-Hill Education.",
        "Redis. (2026). Redis documentation. https://redis.io/docs/latest/",
        "Sommerville, I. (2015). Software engineering (10th ed.). Pearson.",
    ]
    for ref in references:
        p = doc.add_paragraph(style="normal")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(ref)
        set_font(p, 11)
    doc.save(FINAL)
    patch_toc_field(FINAL)
    print(FINAL)


if __name__ == "__main__":
    build_document()
