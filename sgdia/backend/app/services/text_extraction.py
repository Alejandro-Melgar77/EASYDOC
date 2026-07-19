"""Local parsers and OCR for applicant and repository documents."""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass

import docx
import openpyxl
import PyPDF2
import pytesseract
from fastapi import UploadFile
from PIL import Image, ImageEnhance, ImageOps, UnidentifiedImageError

from app.ml.semantic_retriever import LocalSemanticRetriever

try:  # PyMuPDF is optional until the container has been rebuilt with requirements.txt.
    import fitz
except ImportError:  # pragma: no cover - exercised through the unavailable result below.
    fitz = None

logger = logging.getLogger(__name__)

_IMAGE_SUFFIXES = (".png", ".jpg", ".jpeg", ".tiff", ".bmp")
_MAX_OCR_PAGES = 20


@dataclass(frozen=True)
class TextExtractionResult:
    """Result with provenance so users know whether local OCR was involved."""

    text: str
    engine: str
    status: str
    pages_processed: int = 0
    detail: str | None = None


class TextExtractionService:
    """Extract document text locally; never forwards document bytes to an external service."""

    @classmethod
    async def extract_text(cls, file: UploadFile, file_content: bytes) -> str:
        """Compatibility wrapper for existing upload and agent routes."""
        return (await cls.extract_with_metadata(file, file_content)).text

    @classmethod
    async def extract_with_metadata(
        cls, file: UploadFile, file_content: bytes
    ) -> TextExtractionResult:
        filename = (file.filename or "").lower()
        try:
            if filename.endswith(".pdf"):
                return cls._extract_from_pdf(file_content)
            if filename.endswith(".docx"):
                return cls._extracted(cls._extract_from_docx(file_content), "local_docx_parser")
            if filename.endswith(".xlsx"):
                return cls._extracted(cls._extract_from_xlsx(file_content), "local_xlsx_parser")
            if filename.endswith(_IMAGE_SUFFIXES):
                return cls._extract_from_image(file_content)
            if filename.endswith(".txt"):
                return cls._extracted(
                    file_content.decode("utf-8", errors="ignore"), "local_text_parser"
                )
        except (OSError, ValueError, UnidentifiedImageError, PyPDF2.PdfReadError) as error:
            logger.warning("Local extraction failed for %s: %s", filename, error)
            return TextExtractionResult(
                text="",
                engine="local_parser",
                status="no_text",
                detail="No fue posible leer el archivo localmente.",
            )
        except pytesseract.TesseractNotFoundError:
            logger.warning("Tesseract is unavailable for %s", filename)
            return TextExtractionResult(
                text="",
                engine="local_ocr_unavailable",
                status="ocr_unavailable",
                detail="El motor OCR local no esta instalado o no esta disponible.",
            )

        logger.info("Unsupported local extraction type: %s", filename)
        return TextExtractionResult(
            text="",
            engine="local_parser",
            status="unsupported",
            detail="Formato no compatible con la extraccion local.",
        )

    @staticmethod
    def _extracted(text: str, engine: str, pages_processed: int = 0) -> TextExtractionResult:
        cleaned = text.strip()
        return TextExtractionResult(
            text=cleaned,
            engine=engine,
            status="extracted" if cleaned else "no_text",
            pages_processed=pages_processed,
        )

    @classmethod
    def _extract_from_pdf(cls, file_content: bytes) -> TextExtractionResult:
        reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        page_count = len(reader.pages)
        parsed = "\n".join(page.extract_text() or "" for page in reader.pages)
        direct_result = cls._extracted(parsed, "local_pdf_parser", page_count)
        if direct_result.text:
            return direct_result
        return cls._ocr_scanned_pdf(file_content, page_count)

    @classmethod
    def _ocr_scanned_pdf(cls, file_content: bytes, page_count: int) -> TextExtractionResult:
        if fitz is None:
            return TextExtractionResult(
                text="",
                engine="local_ocr_unavailable",
                status="ocr_unavailable",
                pages_processed=0,
                detail="El renderizador local de PDF no esta disponible; reconstruye el servicio backend.",
            )
        document = fitz.open(stream=file_content, filetype="pdf")
        try:
            extracted: list[str] = []
            for page in document[:_MAX_OCR_PAGES]:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                extracted.append(cls._ocr_image(image))
            result = cls._extracted(
                "\n".join(extracted), "local_pdf_ocr", min(page_count, _MAX_OCR_PAGES)
            )
            if page_count > _MAX_OCR_PAGES:
                return TextExtractionResult(
                    text=result.text,
                    engine=result.engine,
                    status=result.status,
                    pages_processed=result.pages_processed,
                    detail=f"OCR local limitado a las primeras {_MAX_OCR_PAGES} paginas.",
                )
            return result
        finally:
            document.close()

    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        document = docx.Document(io.BytesIO(file_content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [
            " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for table in document.tables
            for row in table.rows
        ]
        return "\n".join(part for part in [*paragraphs, *table_cells] if part)

    @staticmethod
    def _extract_from_xlsx(file_content: bytes) -> str:
        workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True, read_only=True)
        text_parts: list[str] = []
        try:
            for sheet in workbook.worksheets:
                text_parts.append(f"Hoja: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join(str(cell) for cell in row if cell is not None)
                    if row_text:
                        text_parts.append(row_text)
        finally:
            workbook.close()
        return "\n".join(text_parts)

    @classmethod
    def _extract_from_image(cls, file_content: bytes) -> TextExtractionResult:
        image = Image.open(io.BytesIO(file_content))
        text = cls._ocr_image(image)
        return cls._extracted(text, "local_image_ocr", pages_processed=1)

    @staticmethod
    def _ocr_image(image: Image.Image) -> str:
        prepared = ImageOps.grayscale(image)
        prepared = ImageEnhance.Contrast(prepared).enhance(1.8)
        return pytesseract.image_to_string(prepared, lang="spa+eng", config="--psm 6")

    @staticmethod
    async def generate_embedding(text: str) -> list[float]:
        """Create a deterministic local semantic vector for repository search."""
        return LocalSemanticRetriever.embed_text(text) if text.strip() else []
