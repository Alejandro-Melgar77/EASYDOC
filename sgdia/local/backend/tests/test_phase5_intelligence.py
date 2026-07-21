"""Focused tests for Phase 5 local intelligence and training artifacts."""

import io
from unittest.mock import patch

import docx
import openpyxl
import pytest
from app.ml.agentic_dataset import build_agentic_cases
from app.ml.demo_dataset import build_demo_dataset
from app.ml.operational_training import build_training_report
from app.ml.semantic_retriever import LocalSemanticRetriever
from app.ml.training_governance import build_training_readiness
from app.services.document_intelligence_service import DocumentIntelligenceService
from app.services.text_extraction import TextExtractionService
from fastapi import UploadFile


def test_document_intelligence_extracts_reviewable_identity_fields() -> None:
    result = DocumentIntelligenceService.analyze_text(
        "Nombre completo: Ana Perez\nCI: 7845123-LP\nCorreo: ana@universidad.edu",
        "documento_identidad.txt",
    )

    assert result.document_type == "identity_document"
    assert result.extraction_status == "extracted"
    assert result.review_required is False
    assert {field.name for field in result.extracted_fields} >= {"full_name", "identity_number"}


def test_agentic_corpus_covers_every_synthetic_policy() -> None:
    corpus = build_agentic_cases()
    codes = {case["policy_code"] for case in corpus["cases"]}

    assert corpus["is_synthetic"] is True
    assert codes == {"CE", "CG", "CN", "HI", "HN", "RB", "RD", "RE", "RT", "TI"}
    assert all(case["expected_action"] == "recommend_policy" for case in corpus["cases"])


def test_training_report_uses_a_held_out_july_validation_period() -> None:
    dataset = build_demo_dataset()
    dataset["agentic_cases"] = build_agentic_cases()["cases"]
    report = build_training_report(dataset)

    assert report["data_provenance"]["is_synthetic"] is True
    assert report["data_provenance"]["training_samples"] > 0
    assert report["data_provenance"]["validation_samples"] > 0
    assert report["validation_period"] == {"from": "2026-07-01", "to": "2026-07-31"}
    assert 0.0 <= report["metrics"]["route_match_rate"] <= 1.0
    assert report["metrics"]["guidance_match_rate"] >= 0.9
    assert report["training_readiness"]["decision"] == "blocked"


def test_local_semantic_retriever_is_deterministic_and_relevant() -> None:
    retriever = LocalSemanticRetriever.from_agentic_cases(build_agentic_cases()["cases"])

    first = retriever.policy_scores("Vengo de otra universidad y deseo homologar materias.")
    second = retriever.policy_scores("Vengo de otra universidad y deseo homologar materias.")

    assert first == second
    assert max(first, key=first.get) == "HN"
    embedding = LocalSemanticRetriever.embed_text("historial academico y homologacion")
    assert len(embedding) == 256
    assert any(embedding)


def test_training_governance_blocks_synthetic_data() -> None:
    readiness = build_training_readiness(build_demo_dataset())

    assert readiness["automation_enabled"] is False
    assert readiness["data_is_synthetic_only"] is True
    assert readiness["decision"] == "blocked"


@pytest.mark.anyio
async def test_local_docx_and_xlsx_extractors_preserve_provenance() -> None:
    docx_buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Nombre completo: Ana Perez")
    document.add_table(rows=1, cols=2).rows[0].cells[0].text = "CI"
    document.tables[0].rows[0].cells[1].text = "7845123-LP"
    document.save(docx_buffer)
    docx_upload = UploadFile(filename="identidad.docx", file=io.BytesIO(docx_buffer.getvalue()))

    docx_result = await TextExtractionService.extract_with_metadata(
        docx_upload, docx_buffer.getvalue()
    )

    assert docx_result.engine == "local_docx_parser"
    assert "Ana Perez" in docx_result.text
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Registro"
    sheet.append(["Registro universitario", "12345678"])
    xlsx_buffer = io.BytesIO()
    workbook.save(xlsx_buffer)
    xlsx_upload = UploadFile(filename="registro.xlsx", file=io.BytesIO(xlsx_buffer.getvalue()))

    xlsx_result = await TextExtractionService.extract_with_metadata(
        xlsx_upload, xlsx_buffer.getvalue()
    )

    assert xlsx_result.engine == "local_xlsx_parser"
    assert "Registro universitario" in xlsx_result.text


@pytest.mark.anyio
async def test_local_image_ocr_marks_its_engine() -> None:
    image_buffer = io.BytesIO()
    from PIL import Image

    Image.new("RGB", (40, 40), "white").save(image_buffer, format="PNG")
    upload = UploadFile(filename="carnet.png", file=io.BytesIO(image_buffer.getvalue()))
    with patch(
        "app.services.text_extraction.pytesseract.image_to_string", return_value="CI: 7845123"
    ):
        result = await TextExtractionService.extract_with_metadata(upload, image_buffer.getvalue())

    assert result.engine == "local_image_ocr"
    assert result.status == "extracted"
    assert result.pages_processed == 1
