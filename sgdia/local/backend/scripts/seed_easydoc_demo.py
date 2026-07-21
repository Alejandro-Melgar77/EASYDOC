#!/usr/bin/env python3
"""Load the deterministic March-July EASYDOC demo dataset into MongoDB.

The loader is idempotent: it updates documents marked ``is_synthetic`` instead
of creating a second copy of policies, workflow instances, or repository rows.
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import sys
from collections import defaultdict
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from bson import ObjectId
from motor.motor_asyncio import AsyncIOMotorDatabase
from pymongo import ASCENDING

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

DATASET_PATH = ROOT / "app" / "ml" / "training_data" / "easydoc_university_march_july.json"
AGENT_CASES_PATH = ROOT / "app" / "ml" / "training_data" / "easydoc_agentic_cases.json"
TRAINING_REPORT_PATH = ROOT / "app" / "ml" / "artifacts" / "easydoc_training_report.json"

LANE_COLORS = ("#2D7180", "#278A89", "#3A78A0", "#8A6C4B", "#7E5C9A", "#547A5B")


def _parse_datetime(value: str) -> datetime:
    return datetime.fromisoformat(value.replace("Z", "+00:00")).astimezone(UTC)


def _form_definition(item: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    return {
        "questions": [
            {
                "id": field["key"],
                "label": field["label"],
                "type": "textarea" if field["key"] in {"justificacion", "motivo"} else "text",
                "required": bool(field.get("required", True)),
                "options": "",
            }
            for field in item["fields"]
        ],
        "attachments": [
            {
                "id": requirement.rsplit(".", 1)[0],
                "label": requirement.replace("_", " ").rsplit(".", 1)[0].title(),
                "acceptedFormats": "application/pdf,image/*,.docx,.xlsx",
                "required": True,
            }
            for requirement in item["requirements"]
        ],
    }


def _diagram(
    policy: dict[str, Any], events: list[dict[str, Any]], staff_ids: dict[str, str]
) -> dict[str, Any]:
    """Build a multi-lane UML 2.5+ activity diagram from the seeded route."""
    stage_metadata = {
        event["stage"]: {
            "department": event["department"],
            "assignee": staff_ids.get(event["worker_id"], ""),
        }
        for event in events
    }
    departments: list[str] = []
    for stage in policy["route"]:
        department = stage_metadata.get(stage, {}).get("department", policy["owner_department"])
        if department not in departments:
            departments.append(department)
    lanes = [
        {"id": f"lane-{index}", "name": department, "color": LANE_COLORS[index % len(LANE_COLORS)]}
        for index, department in enumerate(departments)
    ]
    lane_by_department = {lane["name"]: lane["id"] for lane in lanes}

    def stage_node(stage: str, index: int) -> dict[str, Any]:
        metadata = stage_metadata.get(stage, {})
        department = str(metadata.get("department", policy["owner_department"]))
        lane_index = departments.index(department)
        return {
            "id": f"activity-{index}",
            "type": "activity",
            "x": 210 + (index * 220),
            "y": 74 + (lane_index * 184),
            "label": stage,
            "department": department,
            "assignee": metadata.get("assignee", ""),
            "laneId": lane_by_department[department],
        }

    first_department = departments[0]
    last_department = departments[-1]
    nodes: list[dict[str, Any]] = [
        {
            "id": "start",
            "type": "start",
            "x": 80,
            "y": 88,
            "label": "Inicio",
            "department": first_department,
            "laneId": lane_by_department[first_department],
        }
    ]
    stage_nodes = [stage_node(stage, index) for index, stage in enumerate(policy["route"], start=1)]
    nodes.extend(stage_nodes)
    edges: list[dict[str, Any]] = []

    if policy["code"] == "TI" and len(stage_nodes) >= 5:
        revision = stage_nodes[1]
        parallel_a = stage_nodes[2]
        parallel_b = stage_nodes[3]
        fork = {
            "id": "fork-verificaciones",
            "type": "fork",
            "x": revision["x"] + 170,
            "y": revision["y"] + 26,
            "label": "Fork de verificaciones",
            "department": revision["department"],
            "laneId": revision["laneId"],
        }
        join = {
            "id": "join-verificaciones",
            "type": "join",
            "x": parallel_b["x"] + 175,
            "y": parallel_b["y"] + 26,
            "label": "Join de verificaciones",
            "department": parallel_b["department"],
            "laneId": parallel_b["laneId"],
        }
        nodes.extend([fork, join])
        connections = [
            ("start", stage_nodes[0]["id"]),
            (stage_nodes[0]["id"], revision["id"]),
            (revision["id"], fork["id"]),
            (fork["id"], parallel_a["id"]),
            (fork["id"], parallel_b["id"]),
            (parallel_a["id"], join["id"]),
            (parallel_b["id"], join["id"]),
            (join["id"], stage_nodes[4]["id"]),
        ]
        connections.extend(
            (stage_nodes[index]["id"], stage_nodes[index + 1]["id"])
            for index in range(4, len(stage_nodes) - 1)
        )
    else:
        connections = [("start", stage_nodes[0]["id"])]
        connections.extend(
            (stage_nodes[index]["id"], stage_nodes[index + 1]["id"])
            for index in range(len(stage_nodes) - 1)
        )

    end_x = (stage_nodes[-1]["x"] + 220) if stage_nodes else 300
    nodes.append(
        {
            "id": "end",
            "type": "end",
            "x": end_x,
            "y": 88 + (departments.index(last_department) * 184),
            "label": "Fin",
            "department": last_department,
            "laneId": lane_by_department[last_department],
        }
    )
    connections.append((stage_nodes[-1]["id"], "end"))
    for index, (source, target) in enumerate(connections, start=1):
        edges.append(
            {"id": f"edge-{index}", "from": source, "to": target, "label": "", "kind": "control"}
        )
    return {"nodes": nodes, "edges": edges, "lanes": lanes}


async def _ensure_indexes(database: AsyncIOMotorDatabase) -> None:
    await database["policies"].create_index("demo_code", unique=True, sparse=True)
    await database["policy_versions"].create_index(
        [("policy_id", ASCENDING), ("version", ASCENDING)],
        unique=True,
        name="unique_policy_version",
    )
    await database["workflow_instances"].create_index("demo_workflow_id", unique=True, sparse=True)
    await database["workflow_events"].create_index("demo_event_id", unique=True, sparse=True)
    await database["demo_users"].create_index("source_id", unique=True)
    await database["demo_document_repositories"].create_index("repository_key", unique=True)
    await database["agent_knowledge_cases"].create_index("case_id", unique=True)
    await database["ml_training_runs"].create_index(
        [("report_type", ASCENDING), ("model_name", ASCENDING)], unique=True
    )
    await database["service_requests"].create_index("tracking_code", unique=True)
    await database["workflow_tasks"].create_index(
        [("request_id", ASCENDING), ("node_id", ASCENDING), ("status", ASCENDING)]
    )
    await database["repository_folders"].create_index("repository_key", unique=True, sparse=True)
    await database["policy_artifacts"].create_index(
        [("policy_id", ASCENDING), ("artifact_type", ASCENDING)], unique=True
    )


async def _seed_staff_users(
    database: AsyncIOMotorDatabase, users: list[dict[str, Any]]
) -> dict[str, str]:
    """Materialize the synthetic staff in the real login collection for demo work queues."""
    from app.core.security import hash_password

    role_names = {
        "director": "Director",
        "manager": "Gerente",
        "coordinator": "Coordinador",
        "official": "Funcionario",
        "secretary": "Funcionario",
        "worker": "Trabajador",
    }
    role_documents = {
        role["name"]: role
        async for role in database["roles"].find({"name": {"$in": list(set(role_names.values()))}})
    }
    user_ids: dict[str, str] = {}
    now = datetime.now(UTC)
    for user in users:
        if not user.get("login_required"):
            continue
        role_name = role_names.get(str(user["role"]), "Trabajador")
        role = role_documents.get(role_name)
        update = {
            "name": user["name"],
            "full_name": user["name"],
            "department": user["department"],
            "demo_source_id": user["id"],
            "is_synthetic": True,
            "status": "active",
            "is_active": True,
            "updated_at": now,
        }
        if role is not None:
            update["role_id"] = str(role["_id"])
        await database["users"].update_one(
            {"email": user["email"]},
            {
                "$set": update,
                "$setOnInsert": {
                    "email": user["email"],
                    "password_hash": hash_password("password123"),
                    "created_at": now,
                    "is_deleted": False,
                },
            },
            upsert=True,
        )
        stored = await database["users"].find_one({"email": user["email"]})
        if stored is not None:
            user_ids[str(user["id"])] = str(stored["_id"])
    return user_ids


def _policy_master_document(policy: dict[str, Any], diagram: dict[str, Any]) -> bytes:
    """Create the editable DOCX master used by high-ranking policy editors."""
    from io import BytesIO

    from docx import Document

    document = Document()
    document.add_heading(f"EASYDOC - Politica de negocio: {policy['name']}", level=0)
    document.add_paragraph("Documento maestro sintetico para demostracion local.")
    document.add_heading("Responsabilidad", level=1)
    document.add_paragraph(f"Unidad propietaria: {policy['owner_department']}")
    document.add_heading("Flujo aprobado", level=1)
    for index, stage in enumerate(policy["route"], start=1):
        document.add_paragraph(f"{index}. {stage}", style="List Number")
    document.add_heading("Carriles UML 2.5+", level=1)
    for lane in diagram["lanes"]:
        document.add_paragraph(lane["name"], style="List Bullet")
    document.add_paragraph(
        "Las ediciones deben publicarse desde el diagramador y conservaran su version y bitacora."
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _requirements_matrix(policy: dict[str, Any], form: dict[str, Any]) -> bytes:
    """Create a spreadsheet artifact that documents the dynamic form requirements."""
    from io import BytesIO

    from openpyxl import Workbook

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Requisitos"
    worksheet.append(["Politica", policy["name"]])
    worksheet.append([])
    worksheet.append(["Tipo", "Identificador", "Etiqueta", "Obligatorio", "Formatos"])
    for question in form["questions"]:
        worksheet.append(
            ["Pregunta", question["id"], question["label"], question["required"], question["type"]]
        )
    for attachment in form["attachments"]:
        worksheet.append(
            [
                "Adjunto",
                attachment["id"],
                attachment["label"],
                attachment["required"],
                attachment["acceptedFormats"],
            ]
        )
    worksheet.freeze_panes = "A4"
    worksheet.column_dimensions["A"].width = 16
    worksheet.column_dimensions["B"].width = 28
    worksheet.column_dimensions["C"].width = 42
    worksheet.column_dimensions["D"].width = 14
    worksheet.column_dimensions["E"].width = 38
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


async def _upload_artifact(
    content: bytes, filename: str, content_type: str, folder_path: str
) -> str | None:
    """Store artifacts when MinIO is available without making demo seeding brittle."""
    try:
        from app.services.storage_service import StorageService

        return await StorageService().upload_file_bytes(
            content, filename, content_type, folder_path
        )
    except Exception:
        return None


async def _upsert_policy_artifacts(
    database: AsyncIOMotorDatabase,
    policy: dict[str, Any],
    policy_id: ObjectId,
    diagram: dict[str, Any],
    form: dict[str, Any],
) -> list[dict[str, Any]]:
    """Persist the master document, diagram source and requirements matrix for one policy."""
    artifact_specs = [
        {
            "artifact_type": "master_docx",
            "title": "Politica maestra editable",
            "filename": f"{policy['code']}_politica_maestra.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content": _policy_master_document(policy, diagram),
            "editable_roles": ["Administrador", "Director", "Gerente"],
        },
        {
            "artifact_type": "diagram_json",
            "title": "Diagrama de actividades UML 2.5+",
            "filename": f"{policy['code']}_diagrama_uml.json",
            "content_type": "application/json",
            "content": json.dumps(diagram, ensure_ascii=True, indent=2).encode("utf-8"),
            "editable_roles": ["Administrador", "Director", "Gerente", "Coordinador"],
        },
        {
            "artifact_type": "requirements_xlsx",
            "title": "Matriz de formulario y requisitos",
            "filename": f"{policy['code']}_requisitos.xlsx",
            "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "content": _requirements_matrix(policy, form),
            "editable_roles": ["Administrador", "Director", "Gerente", "Coordinador"],
        },
    ]
    now = datetime.now(UTC)
    artifacts: list[dict[str, Any]] = []
    owner = await database["users"].find_one({"email": "directora@easydoc.edu"})
    owner_id = str(owner["_id"]) if owner else "easydoc-demo-seed"
    for spec in artifact_specs:
        file_key = await _upload_artifact(
            spec["content"],
            spec["filename"],
            spec["content_type"],
            f"policies/{policy['code'].lower()}/",
        )
        source_key = f"{policy['code']}:{spec['artifact_type']}"
        document_payload = {
            "title": spec["title"],
            "filename": spec["filename"],
            "content_type": spec["content_type"],
            "size_bytes": len(spec["content"]),
            "description": f"Artefacto sintetico de la politica {policy['name']}.",
            "tags": ["politica", policy["code"].lower(), "sintetico"],
            "folder_id": f"policy-{policy['code'].lower()}",
            "owner_id": owner_id,
            "file_key": file_key or f"metadata-only/{spec['filename']}",
            "version": 1,
            "is_active": True,
            "permissions": [],
            "policy_id": str(policy_id),
            "artifact_type": spec["artifact_type"],
            "editable_roles": spec["editable_roles"],
            "is_synthetic": True,
            "storage_available": file_key is not None,
            "updated_at": now,
        }
        await database["documents"].update_one(
            {"demo_artifact_key": source_key},
            {
                "$set": document_payload,
                "$setOnInsert": {"demo_artifact_key": source_key, "created_at": now},
            },
            upsert=True,
        )
        document = await database["documents"].find_one({"demo_artifact_key": source_key})
        assert document is not None
        artifact = {
            "id": str(document["_id"]),
            "title": spec["title"],
            "filename": spec["filename"],
            "artifact_type": spec["artifact_type"],
            "document_id": str(document["_id"]),
            "editable_roles": spec["editable_roles"],
            "storage_available": file_key is not None,
        }
        await database["policy_artifacts"].update_one(
            {"policy_id": str(policy_id), "artifact_type": spec["artifact_type"]},
            {
                "$set": {
                    **artifact,
                    "policy_id": str(policy_id),
                    "is_synthetic": True,
                    "updated_at": now,
                }
            },
            upsert=True,
        )
        artifacts.append(artifact)
    return artifacts


def _request_status(status: str) -> str:
    return {
        "recibido": "received",
        "en_curso": "in_progress",
        "observado": "observed",
        "aprobado": "approved",
        "rechazado": "rejected",
        "desechado": "discarded",
        "terminado": "completed",
    }.get(status, "received")


async def _seed_operational_records(
    database: AsyncIOMotorDatabase,
    dataset: dict[str, Any],
    policy_ids: dict[str, ObjectId],
    staff_ids: dict[str, str],
) -> None:
    """Project the synthetic training history into request, task and repository collections."""
    events_by_workflow: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for event in dataset["workflow_events"]:
        events_by_workflow[event["workflow_id"]].append(event)

    now = datetime.now(UTC)
    sensitive_codes = {"CE", "TI", "RD", "RB"}
    for workflow in dataset["workflow_instances"]:
        workflow_events = events_by_workflow[workflow["id"]]
        sequence = int(str(workflow["id"]).rsplit("_", 1)[-1])
        is_active = sequence % 9 == 0
        current_event = workflow_events[min(2, len(workflow_events) - 1)] if is_active else None
        final_status = "in_progress" if is_active else _request_status(str(workflow["status"]))
        receipt_pin = f"{sequence:08d}"
        tracking_code = workflow["request_code"]
        timeline = [
            {
                "at": _parse_datetime(event["finished_at"]),
                "status": _request_status(event["status"]),
                "detail": f"{event['stage']} - {event['department']}",
                "actor_id": staff_ids.get(event["worker_id"]),
            }
            for event in workflow_events
        ]
        attachments = [
            {
                "id": f"{workflow['id']}-{index}",
                "requirement_id": requirement.rsplit(".", 1)[0],
                "filename": f"{tracking_code}_{requirement}",
                "content_type": "application/pdf" if requirement.endswith("pdf") else "image/jpeg",
                "size_bytes": 24_576 + (index * 1_024),
                "uploaded_at": _parse_datetime(workflow["created_at"]),
                "file_key": f"synthetic-requests/{tracking_code}/{requirement}",
                "intelligence": None,
            }
            for index, requirement in enumerate(workflow["requirements"])
        ]
        current_stage = (
            {
                "id": f"activity-{min(3, len(workflow_events))}",
                "label": current_event["stage"],
                "department": current_event["department"],
                "assignee_id": staff_ids.get(current_event["worker_id"]),
            }
            if current_event
            else None
        )
        completed = final_status in {"completed", "rejected", "discarded"}
        response = None
        if completed:
            outcome = {
                "completed": "Solicitud totalmente culminada y registrada en el expediente academico.",
                "rejected": "Solicitud concluida con resultado no aprobado. Revise la respuesta institucional.",
                "discarded": "Solicitud concluida sin tramitacion por falta de condiciones de admisibilidad.",
            }[final_status]
            response = {
                "message": outcome,
                "published_at": _parse_datetime(workflow["completed_at"]),
                "published_by": staff_ids.get(workflow_events[-1]["worker_id"]),
                "requires_approval": workflow["policy_code"] in sensitive_codes,
                "approval_status": "approved",
            }
        request = {
            "demo_workflow_id": workflow["id"],
            "tracking_code": tracking_code,
            "receipt_pin_hash": hashlib.sha256(receipt_pin.encode("utf-8")).hexdigest(),
            "demo_receipt_pin": receipt_pin,
            "policy_id": str(policy_ids[workflow["policy_code"]]),
            "service_title": workflow["policy_name"],
            "applicant": {
                "full_name": workflow["student_name"],
                "email": None,
                "university_id": None,
            },
            "answers": {},
            "attachments": attachments,
            "intake_ready": True,
            "priority": "high" if workflow["priority"] == "alta" else "normal",
            "status": final_status,
            "current_node_id": current_stage["id"] if current_stage else None,
            "current_stage": current_stage,
            "active_node_ids": [current_stage["id"]] if current_stage else [],
            "completed_node_ids": []
            if is_active
            else [f"activity-{index}" for index in range(1, len(workflow_events) + 1)],
            "timeline": timeline,
            "final_response": response,
            "is_fully_completed": completed,
            "is_synthetic": True,
            "created_at": _parse_datetime(workflow["created_at"]),
            "updated_at": _parse_datetime(
                current_event["finished_at"] if current_event else workflow["completed_at"]
            ),
        }
        await database["service_requests"].update_one(
            {"demo_workflow_id": workflow["id"]}, {"$set": request}, upsert=True
        )
        stored_request = await database["service_requests"].find_one(
            {"demo_workflow_id": workflow["id"]}
        )
        assert stored_request is not None
        if current_event is not None:
            await database["workflow_tasks"].update_one(
                {"demo_task_key": f"{workflow['id']}:{current_stage['id']}"},
                {
                    "$set": {
                        "request_id": stored_request["_id"],
                        "tracking_code": tracking_code,
                        "service_title": workflow["policy_name"],
                        "node_id": current_stage["id"],
                        "title": current_stage["label"],
                        "department": current_stage["department"],
                        "assignee_id": current_stage["assignee_id"],
                        "status": "in_progress",
                        "priority": request["priority"],
                        "is_synthetic": True,
                        "created_at": request["created_at"],
                        "updated_at": request["updated_at"],
                    },
                    "$setOnInsert": {"demo_task_key": f"{workflow['id']}:{current_stage['id']}"},
                },
                upsert=True,
            )

    for department, workers in dataset["document_repositories"].items():
        department_key = f"department:{department}"
        await database["repository_folders"].update_one(
            {"repository_key": department_key},
            {
                "$set": {
                    "repository_key": department_key,
                    "name": department,
                    "parent_repository_key": "easydoc:departments",
                    "department": department,
                    "is_synthetic": True,
                    "updated_at": now,
                },
                "$setOnInsert": {"created_at": now},
            },
            upsert=True,
        )
        for worker, files in workers.items():
            worker_key = f"worker:{department}:{worker}"
            await database["repository_folders"].update_one(
                {"repository_key": worker_key},
                {
                    "$set": {
                        "repository_key": worker_key,
                        "name": worker,
                        "parent_repository_key": department_key,
                        "department": department,
                        "worker": worker,
                        "is_synthetic": True,
                        "updated_at": now,
                    },
                    "$setOnInsert": {"created_at": now},
                },
                upsert=True,
            )
            for item in files:
                entry_key = f"{worker_key}:{item['workflow_id']}:{item['filename']}"
                await database["repository_entries"].update_one(
                    {"entry_key": entry_key},
                    {
                        "$set": {
                            **item,
                            "entry_key": entry_key,
                            "repository_key": worker_key,
                            "department": department,
                            "worker": worker,
                            "is_synthetic": True,
                        }
                    },
                    upsert=True,
                )


async def seed_demo(database: AsyncIOMotorDatabase, dataset: dict[str, Any]) -> dict[str, int]:
    """Upsert every synthetic record and return the resulting counters."""
    await _ensure_indexes(database)
    forms = {item["policy_code"]: item for item in dataset["dynamic_forms"]}
    events_by_policy: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for event in dataset["workflow_events"]:
        events_by_policy[event["policy_code"]].append(event)

    staff_ids = await _seed_staff_users(database, dataset["users"])

    policy_ids: dict[str, ObjectId] = {}
    for policy in dataset["business_policies"]:
        first_event = events_by_policy[policy["code"]]
        created_at = _parse_datetime(dataset["period"]["from"])
        updated_at = _parse_datetime(dataset["period"]["to"])
        diagram = _diagram(policy, first_event, staff_ids)
        form_definition = _form_definition(forms[policy["code"]])
        document = {
            "demo_code": policy["code"],
            "title": policy["name"],
            "description": f"Tramite academico de {policy['owner_department']} (datos sinteticos).",
            "diagram_data": diagram,
            "form_definition": form_definition,
            "status": "published",
            "version": 1,
            "created_by": "easydoc-demo-seed",
            "created_at": created_at,
            "updated_at": updated_at,
            "is_deleted": False,
            "is_synthetic": True,
            "dataset_period": dataset["period"],
        }
        await database["policies"].update_one(
            {"demo_code": policy["code"]}, {"$set": document}, upsert=True
        )
        stored = await database["policies"].find_one({"demo_code": policy["code"]})
        assert stored is not None
        policy_ids[policy["code"]] = stored["_id"]
        artifacts = await _upsert_policy_artifacts(
            database, policy, stored["_id"], diagram, form_definition
        )
        await database["policies"].update_one(
            {"_id": stored["_id"]}, {"$set": {"artifacts": artifacts, "updated_at": updated_at}}
        )
        document["artifacts"] = artifacts
        await database["policy_versions"].update_one(
            {"policy_id": str(stored["_id"]), "version": 1},
            {
                "$set": {
                    **document,
                    "policy_id": str(stored["_id"]),
                    "change_summary": "Synthetic EASYDOC baseline",
                }
            },
            upsert=True,
        )

    for user in dataset["users"]:
        actual_user_id = staff_ids.get(str(user["id"]))
        await database["demo_users"].update_one(
            {"source_id": user["id"]},
            {
                "$set": {
                    **user,
                    "actual_user_id": actual_user_id,
                    "is_synthetic": True,
                }
            },
            upsert=True,
        )

    if AGENT_CASES_PATH.exists():
        agent_cases = json.loads(AGENT_CASES_PATH.read_text(encoding="utf-8"))
        for case in agent_cases.get("cases", []):
            await database["agent_knowledge_cases"].update_one(
                {"case_id": case["id"]},
                {"$set": {**case, "case_id": case["id"], "is_synthetic": True}},
                upsert=True,
            )

    if TRAINING_REPORT_PATH.exists():
        report = json.loads(TRAINING_REPORT_PATH.read_text(encoding="utf-8"))
        await database["ml_training_runs"].update_one(
            {"report_type": report["report_type"], "model_name": report["model_name"]},
            {"$set": {**report, "is_synthetic": True}},
            upsert=True,
        )

    for workflow in dataset["workflow_instances"]:
        code = workflow["policy_code"]
        document = {
            **workflow,
            "demo_workflow_id": workflow["id"],
            "policy_id": str(policy_ids[code]),
            "started_by": "guest-demo",
            "history": [],
            "is_synthetic": True,
            "created_at": _parse_datetime(workflow["created_at"]),
            "updated_at": _parse_datetime(workflow.get("completed_at", workflow["created_at"])),
        }
        await database["workflow_instances"].update_one(
            {"demo_workflow_id": workflow["id"]}, {"$set": document}, upsert=True
        )

    for index, event in enumerate(dataset["workflow_events"]):
        document = {
            **event,
            "demo_event_id": f"{event['workflow_id']}:{index}",
            "policy_id": str(policy_ids[event["policy_code"]]),
            "actor_id": staff_ids.get(event["worker_id"]),
            "is_synthetic": True,
            "at": _parse_datetime(event["finished_at"]),
        }
        await database["workflow_events"].update_one(
            {"demo_event_id": document["demo_event_id"]}, {"$set": document}, upsert=True
        )

    for department, workers in dataset["document_repositories"].items():
        for worker, files in workers.items():
            repository_key = f"{department}:{worker}"
            await database["demo_document_repositories"].update_one(
                {"repository_key": repository_key},
                {
                    "$set": {
                        "repository_key": repository_key,
                        "department": department,
                        "worker": worker,
                        "files": files,
                        "is_synthetic": True,
                        "period": dataset["period"],
                    }
                },
                upsert=True,
            )

    await _seed_operational_records(database, dataset, policy_ids, staff_ids)

    return {
        "policies": await database["policies"].count_documents({"is_synthetic": True}),
        "workflow_instances": await database["workflow_instances"].count_documents(
            {"is_synthetic": True}
        ),
        "workflow_events": await database["workflow_events"].count_documents(
            {"is_synthetic": True}
        ),
        "repositories": await database["demo_document_repositories"].count_documents(
            {"is_synthetic": True}
        ),
        "agent_cases": await database["agent_knowledge_cases"].count_documents(
            {"is_synthetic": True}
        ),
        "training_runs": await database["ml_training_runs"].count_documents({"is_synthetic": True}),
        "service_requests": await database["service_requests"].count_documents(
            {"is_synthetic": True}
        ),
        "repository_entries": await database["repository_entries"].count_documents(
            {"is_synthetic": True}
        ),
    }


async def main() -> None:
    from app.core.config import get_settings
    from motor.motor_asyncio import AsyncIOMotorClient

    dataset = json.loads(DATASET_PATH.read_text(encoding="utf-8"))
    settings = get_settings()
    client = AsyncIOMotorClient(settings.MONGODB_URL)
    try:
        result = await seed_demo(client[settings.MONGODB_DB_NAME], dataset)
        print(json.dumps(result, indent=2))
    finally:
        client.close()


if __name__ == "__main__":
    asyncio.run(main())
